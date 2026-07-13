from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Cover Page ─────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LG Smart Factory')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0xA5, 0x00, 0x34)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI-Powered Operational Intelligence Platform')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x7A, 0x00, 0x26)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Project Report')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('June 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Helper ──────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xA5, 0x00, 0x34)

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x7A, 0x00, 0x26)

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xA5, 0x00, 0x34)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def para(text):
    doc.add_paragraph(text)

def table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

# ── 1. Project Overview ────────────────────────────────────────────────
h1('1. Project Overview')
para('The LG Smart Factory platform is an enterprise-grade operational intelligence system built for modern manufacturing environments. It integrates AI-powered analytics, real-time monitoring, predictive maintenance, incident management, vision analysis, digital twin simulation, and executive reporting into a unified dashboard.')
para('The platform serves seven user roles across five factory departments — Production, Warehouse, Maintenance, Quality, and Safety — with role-scoped access, live cloud database, and AI-assisted decision-making powered by Google Gemini.')

# ── 2. Architecture ─────────────────────────────────────────────────────
h1('2. System Architecture')

h2('2.1 Technology Stack')
table(['Layer', 'Technology', 'Purpose'], [
    ['Frontend', 'Streamlit', 'Single-page UI with role-based navigation'],
    ['AI / LLM', 'Google Gemini 2.0 Flash / 2.5 Flash', 'Copilot, vision, reports, alert actions'],
    ['Rule Engine', 'Custom Python (SEV levels)', 'Risk detection, predictive scoring, agents'],
    ['Database', 'Supabase (PostgreSQL)', 'Cloud data storage, user auth'],
    ['PDF Export', 'fpdf2', 'Executive report generation'],
    ['Deployment', 'Streamlit Community Cloud', 'Auto-deploy from GitHub'],
])

h2('2.2 Authentication & Roles')
para('SHA-256 hashed password authentication against Supabase users table. Seven roles with scoped page access:')
table(['Role', 'Access Scope'], [
    ['Admin', 'All pages, data entry, reports'],
    ['Factory Manager', 'Dashboard, copilot, executive reports, simulator, agents'],
    ['Production', 'Production page, incident reporting'],
    ['Warehouse', 'Warehouse page, incident reporting'],
    ['Maintenance', 'Maintenance page, predictive analysis'],
    ['Quality', 'Quality page, incident reporting'],
    ['Safety', 'Safety page, vision analysis, incident reporting'],
])

# ── 3. Modules ──────────────────────────────────────────────────────────
h1('3. Module Description')

h2('3.1 Dashboard')
para('Centralized operational view with:')
bullet('KPI cards for all 5 departments (production output, defect rate, downtime, incidents, etc.)')
bullet('Auto-refresh every 30 seconds for live data')
bullet('Task queue and AI recommendations per role')
bullet('Critical alerts panel with severity-coded warnings')
bullet('Breakdown risk detection — identifies machines with "Breakdown Risk" status and suggests actions')

h2('3.2 AI Factory Copilot')
para('Located in modules/llm_engine.py. A Gemini 2.5 Flash-powered conversational assistant that can:')
bullet('Analyze production output, defect trends, and quality metrics')
bullet('Assess maintenance schedules and predict failure risks')
bullet('Review warehouse inventory levels and bottlenecks')
bullet('Evaluate safety incident patterns')
bullet('Generate cross-departmental optimization strategies')

h2('3.3 Predictive Maintenance')
para('Located in modules/predictive_engine.py. Rule-based risk scoring engine that analyzes:')
bullet('Temperature trends')
bullet('Vibration patterns')
bullet('Runtime hours')
bullet('Pressure readings')
bullet('Generates risk levels (Low / Medium / High / Critical) with specific recommendations')

h2('3.4 Visual Incident Reporting')
para('Located in modules/incident_reporting.py. Allows operators to:')
bullet('Upload incident images')
bullet('Tag severity (Low / Medium / High / Critical)')
bullet('Categorize by department')
bullet('View incident feed with filtering and statistics')

h2('3.5 Gemini Vision Analysis')
para('Located in modules/vision_engine.py. Uses Gemini 2.0 Flash to analyze uploaded incident images:')
bullet('Identifies probable root cause')
bullet('Assesses severity level')
bullet('Recommends corrective actions')
bullet('Provides detailed analysis report')

h2('3.6 Agentic AI Operations')
para('Located in modules/agents.py. Five specialized rule-based agents coordinated by a central orchestrator:')
bullet('Production Agent — throughput, defects, OEE')
bullet('Warehouse Agent — inventory levels, stockouts')
bullet('Maintenance Agent — downtime, repair urgency')
bullet('Quality Agent — defect rates, pass/fail trends')
bullet('Safety Agent — incident frequency, risk scoring')
bullet('Coordinator — aggregates findings, prioritizes actions, generates summary')

h2('3.7 Digital Twin Simulator')
para('Located in modules/simulator.py. Scenario-based simulation engine:')
bullet('7 predefined scenarios: demand surge, machine failure, material shortage, power outage, quality crisis, labor shortage, logistics delay')
bullet('Cost impact projection')
bullet('Cascade effect analysis across departments')
bullet('Recovery time estimation')
bullet('Resilience scoring')
bullet('AI-generated strategic response using Gemini')

h2('3.8 Executive Reports')
para('Located in modules/report_generator.py. Generates:')
bullet('Daily / Weekly / Monthly consolidated reports with KPIs from all departments')
bullet('Individual module reports (Production, Warehouse, etc.)')
bullet('Gemini 2.0 Flash narrative analysis')
bullet('PDF export via fpdf2 with LG-branded header (red #a50034)')
bullet('Automatic fallback report generation when Gemini quota is exhausted')

h2('3.9 Breakdown Alerts')
para('Located in modules/breakdown_alerts.py. Real-time detection of:')
bullet('Machine_Status == "Breakdown Risk" from production data')
bullet('Gemini-generated immediate action steps')
bullet('Fallback rule-based actions when API quota is exhausted')
bullet('Visual alert cards using Streamlit native components (error/warning banners)')

# ── 4. Database ─────────────────────────────────────────────────────────
h1('4. Supabase Database')

para('The platform uses Supabase (PostgreSQL) with 7 tables:')

table(['Table', 'Columns', 'Records'], [
    ['production', '8 columns (line, product, output, defects, status, etc.)', '100'],
    ['warehouse', '6 columns (zone, item, stock, capacity, etc.)', '100'],
    ['maintenance', '5 columns (equipment, type, status, etc.)', '100'],
    ['quality', '5 columns (batch, param, value, result, etc.)', '100'],
    ['safety', '5 columns (area, type, severity, etc.)', '100'],
    ['incident_log', '7 columns (reporter, category, image, description, etc.)', 'Variable'],
    ['users', '5 columns (username, password_hash, role, etc.)', '7'],
])

para('Data is seeded from Excel files in datalg2/ via scripts/seed_supabase.py. The modules/database.py helper provides load_table(), insert_record(), and Supabase client initialization.')

# ── 5. Design System ───────────────────────────────────────────────────
h1('5. Design System')

para('The user interface follows LG brand guidelines:')
bullet('Primary color: #a50034 (LG Red)')
bullet('Dark accent: #7a0026')
bullet('Highlight: #d90452')
bullet('Background cream: #faf6f0, #f3ede4, #f0e8dc')
bullet('All UI cards use white backgrounds with subtle shadows')
bullet('Login page uses the same warm cream palette (migrated from dark theme)')
bullet('Logo displayed in sidebar and login page')

# ── 6. Key Technical Decisions ──────────────────────────────────────────
h1('6. Key Technical Decisions')
bullet('Google Gemini exclusively used — no OpenAI/GPT dependencies throughout the platform')
bullet('Fallback generators for reports and alerts when Gemini daily free quota is exhausted (429 errors)')
bullet('fpdf2 for PDF generation with Helvetica font; emoji characters stripped before export to avoid Unicode encoding errors')
bullet('SHA-256 hashed passwords stored in Supabase users table')
bullet('Dashboard panels render as single HTML string per st.markdown() call to prevent orphaned </div> tags appearing as visible text')
bullet('All AI agents use rule-based logic (no LLM calls) for speed and reliability')
bullet('Column casing normalization in load_table() maps database lowercase column names to PascalCase for display')
bullet('Python 3.12 pinned for cloud deployment compatibility')

# ── 7. File Structure ──────────────────────────────────────────────────
h1('7. File Structure')
para('lg-smart-factory-operations/')
bullet('app1.py — Main dashboard (auth, routing, data entry, live incident feed, reports, alerts)')
bullet('requirements.txt — Python dependencies')
bullet('pyproject.toml — Project metadata')
bullet('.streamlit/secrets.toml — API keys (GEMINI_API_KEY, SUPABASE_URL, SUPABASE_KEY)')
bullet('modules/ — 18 Python modules (theme, auth, roles, database, llm_engine, ai_engine, predictive_engine, agents, simulator, incident_reporting, vision_engine, report_generator, breakdown_alerts, production, warehouse, maintenance, quality, safety)')
bullet('datalg2/ — 5 Excel data files')
bullet('scripts/ — seed_supabase.py, seed_users.py')
bullet('data/ — incidents.json, incident_images/')

# ── 8. Dependencies ────────────────────────────────────────────────────
h1('8. Python Dependencies')
table(['Package', 'Version', 'Purpose'], [
    ['streamlit', '1.58.0', 'Web UI framework'],
    ['streamlit-autorefresh', '1.0.1', 'Live dashboard auto-refresh'],
    ['google-generativeai', '0.8.6', 'Gemini AI API'],
    ['supabase', '2.31.0', 'Cloud database client'],
    ['fpdf2', '2.8.7', 'PDF report generation'],
    ['pandas', '3.0.3', 'Data manipulation'],
    ['numpy', '2.4.6', 'Numerical computations'],
    ['Pillow', '12.2.0', 'Image processing'],
    ['openpyxl', '3.1.5', 'Excel file I/O'],
    ['plotly', '(latest)', 'Interactive charts and graphs'],
])

# ── 9. User Credentials ────────────────────────────────────────────────
h1('9. Default User Credentials')
table(['Role', 'Username', 'Password'], [
    ['Admin', 'admin', 'admin@2024'],
    ['Factory Manager', 'manager', 'factory@2024'],
    ['Production', 'production', 'prod@2024'],
    ['Maintenance', 'maintenance', 'maint@2024'],
    ['Warehouse', 'warehouse', 'wh@2024'],
    ['Quality', 'quality', 'qual@2024'],
    ['Safety', 'safety', 'safety@2024'],
])

# ── 10. Setup Instructions ─────────────────────────────────────────────
h1('10. Local Setup Instructions')
para('1. Clone the repository')
para('2. Install dependencies: pip install -r requirements.txt')
para('3. Create .streamlit/secrets.toml with:')
para('    GEMINI_API_KEY = "your-key-here"')
para('    SUPABASE_URL = "https://your-project.supabase.co"')
para('    SUPABASE_KEY = "your-key-here"')
para('4. Seed the database: py scripts/seed_supabase.py and py scripts/seed_users.py')
para('5. Run: py -m streamlit run app1.py')

# ── 11. Deployment ─────────────────────────────────────────────────────
h1('11. Deployment')
para('The application is deployed on Streamlit Community Cloud and automatically redeploys when changes are pushed to the main branch.')
para('Live URL: https://lg-smart-factory-operations.streamlit.app')

# ── 12. Future Enhancements ─────────────────────────────────────────────
h1('12. Future Enhancements')
bullet('Session expiry and auto-logout for enhanced security')
bullet('Auto-redirect to dashboard after successful login')
bullet('Gemini copilot connected to live Supabase data (currently rule-based)')
bullet('Real-time WebSocket updates for dashboard')
bullet('Upgrade to paid Gemini tier for higher API quota')
bullet('Mobile-responsive layout')
bullet('Multi-language support')
bullet('Export reports to Excel in addition to PDF')

# ── Save ────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'LG_Smart_Factory_Project_Report.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
